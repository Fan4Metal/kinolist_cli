from docx import Document
from docx.shared import Cm, Pt, RGBColor


def prepare_document(document):
    top_margin = 1.5
    bottom_margin = 1.5
    left_margin = 2.5
    right_margin = 1.5
    document._body.clear_content()
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(top_margin)
        section.bottom_margin = Cm(bottom_margin)
        section.left_margin = Cm(left_margin)
        section.right_margin = Cm(right_margin)

def add_table(document):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.enum.table import WD_ROW_HEIGHT_RULE

    def set_cell_color(cell, color):
        cell_xml_element = cell._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'), color)
        table_cell_properties.append(shade_obj)

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).width = Cm(7.5)
    table.cell(0, 1).width = Cm(9.38)
    table.style = 'Table Grid'
    table.cell(0, 0).merge(table.cell(1, 0))
    set_cell_color(table.cell(0, 1), 'ffc000')
    table.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[0].height = Cm(0.45)
    table.rows[1].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    table.rows[1].height = Cm(10.92)
    table.cell(0, 0).paragraphs[0].paragraph_format.keep_with_next = True
    table.cell(0, 0).paragraphs[0].paragraph_format.keep_together = True
    document.add_paragraph()

def main():
    document = Document()
    prepare_document(document)
    add_table(document)
    add_table(document)
    add_table(document)
    document.save('test.docx')

if __name__ == '__main__':
    main()
    print('Файл создан')
